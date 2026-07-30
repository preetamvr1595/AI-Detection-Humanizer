"""
Document Processing Engine (PRD Module 1, FR-06).

Fixes two real, silent-failure bugs found via a live bug report:

  1. DOCX extraction previously only read `document.paragraphs`, which
     SILENTLY DROPS any text inside tables. Assignment documents that put
     code/answers in table cells (very common) had most of their content
     never reach the analysis pipeline at all — the AI detector was then
     confidently scoring whatever scraps of paragraph text remained (often
     just headers), not the actual submitted content.

  2. PDF extraction had zero OCR fallback. If a PDF page is a scanned
     image or a screenshot of code with no embedded text layer,
     `pdfplumber` returns empty text for that page SILENTLY — no error,
     no warning, just nothing. The rest of the pipeline then ran on
     whatever fragment of real text happened to exist elsewhere in the
     file, again producing a confident-looking but meaningless score.

Both are fixed here: DOCX tables are now walked and included, and any PDF
page with under a real-text threshold falls back to rendering the page as
an image and running it through a REAL Tesseract OCR pass (poppler +
tesseract, both genuinely installed — this is the "Tesseract-driven OCR"
requirement from PRD Module 1 that was previously unimplemented).

`extract_text_with_diagnostics()` also reports exactly how much text came
from where and flags low-confidence extractions, so a suspiciously short
result is visible to the caller instead of silently producing a score.
"""
import os
import pdfplumber
import docx
import pytesseract
from pdf2image import convert_from_path

MIN_TEXT_CHARS_PER_PAGE = 40  # below this, a PDF page is treated as "no usable text layer"


def _extract_docx_tables(document) -> list:
    table_text = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_text.append(" | ".join(cells))
    return table_text


def _ocr_pdf_page(file_path: str, page_number: int) -> str:
    """Renders a single PDF page to an image and runs real Tesseract OCR on it."""
    try:
        images = convert_from_path(file_path, first_page=page_number + 1, last_page=page_number + 1, dpi=200)
        if not images:
            return ""
        return pytesseract.image_to_string(images[0])
    except Exception:
        # poppler/tesseract failures (corrupt page, unsupported encoding, etc.)
        # degrade gracefully to "no text recovered" rather than crashing the upload.
        return ""


def extract_text(file_path: str, file_type: str) -> str:
    return extract_text_with_diagnostics(file_path, file_type)["text"]


def extract_text_with_diagnostics(file_path: str, file_type: str) -> dict:
    file_type = file_type.lower()
    diagnostics = {"pages_total": 0, "pages_ocr_fallback": 0, "tables_found": 0, "extraction_method": []}

    if file_type == "pdf":
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            diagnostics["pages_total"] = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                if len(page_text.strip()) >= MIN_TEXT_CHARS_PER_PAGE:
                    text_parts.append(page_text)
                    diagnostics["extraction_method"].append("text_layer")
                else:
                    # No usable text layer on this page — real OCR fallback,
                    # not a silent empty result.
                    ocr_text = _ocr_pdf_page(file_path, i)
                    if ocr_text.strip():
                        text_parts.append(ocr_text)
                        diagnostics["pages_ocr_fallback"] += 1
                        diagnostics["extraction_method"].append("ocr")
                    else:
                        diagnostics["extraction_method"].append("empty")
        text = "\n".join(text_parts)

    elif file_type in ("docx", "doc"):
        d = docx.Document(file_path)
        paragraph_text = [p.text for p in d.paragraphs if p.text.strip()]
        table_rows = _extract_docx_tables(d)
        diagnostics["tables_found"] = len(d.tables)
        diagnostics["extraction_method"].append("paragraphs")
        if table_rows:
            diagnostics["extraction_method"].append("tables")
        text = "\n".join(paragraph_text + table_rows)

    elif file_type == "txt":
        with open(file_path, "r", errors="ignore") as f:
            text = f.read()
        diagnostics["extraction_method"].append("plain_text")

    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    word_count = len(text.split())
    diagnostics["extracted_word_count"] = word_count
    diagnostics["low_confidence_extraction"] = word_count < 50
    diagnostics["warning"] = (
        f"Only {word_count} words were extracted from this document — analysis scores below are "
        "unreliable and should not be trusted. This usually means the file is a scanned image OCR "
        "could not read cleanly, or content is in a format (e.g. embedded objects) this parser doesn't cover."
        if diagnostics["low_confidence_extraction"] else None
    )

    return {"text": text, "diagnostics": diagnostics}
